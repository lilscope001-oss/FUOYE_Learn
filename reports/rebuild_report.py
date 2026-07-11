from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "reports" / "FUOYE_Learn_Project_Report.docx"
DFD1 = ROOT / "reports" / "assets" / "dfd_level_1.png"
DFD2 = ROOT / "reports" / "assets" / "dfd_level_2.png"


def style_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def para(doc, text, size=11, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if align:
        p.alignment = align
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        style_run(r, size=16 if level == 1 else 13, bold=True, color=RGBColor(46, 116, 181))
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r, size=10.5)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r, size=10.5)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Description"
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(1))

    para(doc, "PROJECT DEFENSE REPORT", size=12, bold=True, color=RGBColor(47, 142, 99), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "FUOYE Learn", size=28, bold=True, color=RGBColor(35, 75, 61), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "A Web-Based Computer Science Learning and Quiz Management Portal", size=14, color=RGBColor(96, 118, 109), align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Prepared for project defense presentation", size=11, color=RGBColor(96, 118, 109), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    heading(doc, "Project Overview")
    para(doc, "FUOYE Learn is a web-based educational platform designed to help Computer Science students access course quizzes, track learning progress, earn badges, and compare performance through a leaderboard.")
    add_table(doc, [
        ("System Name", "FUOYE Learn"),
        ("System Category", "E-learning and quiz management portal"),
        ("Main Users", "Students and administrator"),
        ("Frontend", "HTML, CSS and JavaScript"),
        ("Backend Service", "Supabase Authentication and Supabase database"),
        ("Deployment Target", "Vercel static hosting"),
    ])

    heading(doc, "Problem Statement")
    para(doc, "Students need a structured way to revise course concepts, test understanding, and monitor learning progress across multiple academic levels. Manual practice does not provide immediate feedback, progression control, XP tracking, or leaderboard comparison.")

    heading(doc, "Aim and Objectives")
    para(doc, "The aim of this project is to develop an interactive learning portal that supports structured quiz practice and progress tracking for FUOYE Computer Science students.")
    bullets(doc, [
        "Provide secure registration and login using Supabase Authentication.",
        "Present Computer Science courses from 100L to 400L in a searchable catalog.",
        "Generate professional timed quiz questions from course content.",
        "Prevent users from moving to the next course until they pass the current quiz.",
        "Record XP, badges, passed courses and learning progress.",
        "Provide leaderboard ranking and admin record management.",
        "Support responsive layouts and dark mode.",
    ])

    heading(doc, "System Architecture")
    para(doc, "The system is implemented as a static client-side web application. HTML pages provide structure, CSS handles design and responsiveness, JavaScript controls interaction and business logic, and Supabase provides authentication and user profile storage.")
    add_table(doc, [
        ("Presentation Layer", "Dashboard, courses, quiz, progress, achievements, leaderboard, profile and settings pages."),
        ("Client Logic Layer", "Authentication checks, quiz generation, scoring, course progression, dark mode, and rendering of dynamic data."),
        ("Backend Service", "Supabase Auth and users/profile database table."),
        ("Local Storage", "Stores passed-course progression and local theme preference."),
    ])

    heading(doc, "Level 1 Data Flow Diagram")
    para(doc, "The Level 1 DFD decomposes FUOYE Learn into major processes: authentication and profile, course catalog, quiz engine, progress and achievements, leaderboard, and admin management.")
    if DFD1.exists():
        doc.add_picture(str(DFD1), width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "Figure 1: Level 1 DFD for FUOYE Learn.", size=9.5, color=RGBColor(96, 118, 109), align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "Level 2 Data Flow Diagram")
    para(doc, "The Level 2 DFD expands the quiz and progress flow. It shows course selection, prerequisite checking, timed quiz generation, answer submission, scoring, XP update, passed-course storage, and the next quiz or retake prompt.")
    if DFD2.exists():
        doc.add_picture(str(DFD2), width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, "Figure 2: Level 2 DFD for quiz, progress and course advancement.", size=9.5, color=RGBColor(96, 118, 109), align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "Major System Modules")
    add_table(doc, [
        ("Authentication", "Signup, login, logout, session checks and profile creation."),
        ("Dashboard", "Displays XP, level, badges, progress, featured courses, activities and top students."),
        ("Courses", "Lists and filters courses by level, semester and search term."),
        ("Quiz", "Generates timed questions, calculates scores, awards XP and controls course advancement."),
        ("Progress", "Tracks completed courses, XP, level and badges."),
        ("Achievements", "Displays locked and unlocked badge milestones."),
        ("Leaderboard", "Ranks users by XP."),
        ("Admin", "Allows authorized management of student records."),
    ])

    heading(doc, "Database and Security Design")
    bullets(doc, [
        "Supabase Authentication manages user identity and login sessions.",
        "The users table stores names, matric numbers, levels, departments, XP, streaks, badges and dark mode preference.",
        "Row Level Security allows users to update their own profiles while the admin email can manage user records.",
        "Passed-course progression is stored locally using a user-specific browser storage key.",
    ])

    heading(doc, "Testing and Deployment Plan")
    numbered(doc, [
        "Run the Supabase SQL setup script.",
        "Confirm signup, login and profile creation.",
        "Test quiz pass/fail progression and XP updates.",
        "Test dashboard, courses, progress, achievements, leaderboard, profile and settings responsiveness.",
        "Deploy the static project to Vercel and retest Supabase connectivity.",
    ])

    heading(doc, "Conclusion")
    para(doc, "FUOYE Learn demonstrates a practical web-based learning system for structured quiz practice, course progression, achievement tracking and leaderboard motivation. It is suitable for defense because it combines clear user workflows, backend authentication, responsive design and measurable learning outcomes.")
    heading(doc, "Defense Talking Points")
    bullets(doc, [
        "The project solves a student revision and learning-tracking problem.",
        "The DFDs explain how data moves from login to course selection, quiz scoring and progress update.",
        "Supabase provides authentication and protected profile storage.",
        "The quiz engine enforces progression by requiring users to pass before moving forward.",
        "The application is responsive, deployable to Vercel and supports dark mode.",
    ])

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    build()
